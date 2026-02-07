"""
Academic Reference Manager

A command-line tool for managing academic references and generating citations
in various styles including Harvard, APA, and IEEE.
"""
import json
import logging
import os
import re
import string
import time
import random
import requests
import unicodedata
from datetime import datetime, timedelta
from difflib import SequenceMatcher
from functools import wraps
from pathlib import Path
from typing import List, Dict, Set, Optional, Tuple, Any, Union, cast, Callable, TypeVar, Type

from docx import Document
from dataclasses import asdict, is_dataclass

# Import ReferenceManager for consolidation
try:
    from src.reference_manager import ReferenceManager
except ImportError:
    try:
        # Fallback for relative import if src is not package root
        from ..reference_manager import ReferenceManager
    except ImportError:
        # Fallback for when running script directly/path issues
        import sys
        sys.path.append(str(Path(__file__).parent.parent))
        from reference_manager import ReferenceManager

import threading

# Singleton manager instance
_manager = None
_manager_lock = threading.Lock()

def _get_manager():
    global _manager
    if _manager is None:
        with _manager_lock:
            if _manager is None:
                _manager = ReferenceManager()
    return _manager

# API Configuration
COLLEGE_CATALOG_API_URL = "http://127.0.0.1:8000"  # URL for the local college catalog API

# Type variable for generic function typing
F = TypeVar('F', bound=Callable[..., Any])

def handle_api_errors(max_retries: int = 3, backoff_factor: float = 0.5):
    """
    Decorator to handle API errors with retries and exponential backoff.
    
    Args:
        max_retries: Maximum number of retry attempts
        backoff_factor: Factor for exponential backoff between retries
    """
    def decorator(func: F) -> F:
        @wraps(func)
        def wrapper(*args, **kwargs):
            last_exception = None
            
            for attempt in range(max_retries + 1):
                try:
                    return func(*args, **kwargs)
                    
                except requests.exceptions.RequestException as e:
                    last_exception = e
                    status_code = getattr(e.response, 'status_code', None)
                    
                    # Don't retry on client errors (4xx) except 429 (Too Many Requests)
                    if status_code and 400 <= status_code < 500 and status_code != 429:
                        logger.error(f"API client error {status_code}: {str(e)}")
                        break
                        
                    # Handle rate limiting
                    if status_code == 429:
                        retry_after = int(e.response.headers.get('Retry-After', 5))
                        logger.warning(f"Rate limited. Retrying after {retry_after} seconds...")
                        time.sleep(retry_after)
                        continue
                        
                    # Calculate backoff time with jitter
                    backoff = min(backoff_factor * (2 ** attempt) + random.uniform(0, 1), 60)  # Max 60 seconds
                    
                    if attempt < max_retries:
                        logger.warning(f"API request failed (attempt {attempt + 1}/{max_retries}): {str(e)}")
                        logger.info(f"Retrying in {backoff:.2f} seconds...")
                        time.sleep(backoff)
                    
                except Exception as e:
                    last_exception = e
                    logger.error(f"Unexpected error in {func.__name__}: {str(e)}")
                    if attempt >= max_retries:
                        break
                    
                    backoff = min(backoff_factor * (2 ** attempt) + random.uniform(0, 1), 60)
                    logger.info(f"Retrying in {backoff:.2f} seconds...")
                    time.sleep(backoff)
            
            # If we got here, all retries failed
            error_message = f"Failed after {max_retries + 1} attempts"
            if last_exception:
                error_message += f": {str(last_exception)}"
            logger.error(error_message)
            return None
            
        return cast(F, wrapper)
    return decorator

# Define cache file path
CACHE_DIR = Path.home() / ".cache" / "referencing"
CACHE_DIR.mkdir(parents=True, exist_ok=True)
CACHE_FILE = CACHE_DIR / "references_cache.json"

# Define download folder and word filename
DOWNLOAD_FOLDER = Path.home() / "Downloads"
WORD_FILENAME = "references.docx"

# Application configuration and utilities
from .config_loader import Config
from .utils.api_utils import (
    safe_crossref_request,
    safe_google_books_request,
    APIError
)

# Input validation
try:
    from referencing.utils.input_validation import (
        validate_menu_choice,
        validate_author_name,
        validate_search_query,
        get_valid_input,
        confirm_action
    )
except ImportError:
    # Fallback to older function names if available
    try:
        from referencing.utils.input_validation import (
            get_valid_menu_choice as validate_menu_choice,
            get_non_empty_input as validate_author_name,
            get_non_empty_input as validate_search_query,
            get_non_empty_input as get_valid_input,
            get_user_confirmation as confirm_action,
        )
    except ImportError:
        # Fallback implementations if input_validation is not available
        def validate_menu_choice(*args, **kwargs):
            return input("Enter your choice: ")
            
        def validate_author_name(*args, **kwargs):
            return input("Enter author name: ").strip()
            
        def validate_search_query(*args, **kwargs):
            return input("Enter search query: ").strip()
            
        def get_valid_input(*args, **kwargs):
            return input().strip()
            
        def confirm_action(*args, **kwargs):
            return input("Continue? (y/n): ").lower().startswith('y')

# Set up logging
logging.basicConfig(
    level=getattr(logging, Config.LOG_LEVEL, logging.INFO),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(Config.LOG_FILE),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Type aliases for better code readability
Publication = Dict[str, Any]
Author = Dict[str, str]
CacheDict = Dict[str, Any]

# Compatibility note:
# This file tries to import modern validation helper names (e.g. validate_menu_choice,
# validate_author_name, validate_search_query). Older versions of the project
# expose helper functions under different names in `input_validation.py` (for
# example: get_valid_menu_choice, get_user_confirmation, get_non_empty_input).
# The import logic falls back to those older names so the CLI and tests continue
# to work if the module uses either naming convention. If you refactor
# `input_validation.py`, keep one of these sets of names or update the imports
# here accordingly.

########################################
# CACHE MANAGEMENT
########################################

########################################
# CACHE MANAGEMENT (Wrappers)
########################################

def load_cache() -> Dict:
    """Load cache via ReferenceManager."""
    return _get_manager().cache

def save_cache(cache: Dict) -> None:
    """Save cache via ReferenceManager."""
    manager = _get_manager()
    manager.cache = cache
    manager._save_cache()

########################################
# RESULT RANKING
########################################

def rank_results(results: List[Dict[str, Any]], query: str) -> List[Dict[str, Any]]:
    """Rank search results by relevance."""
    if not results or not query:
        return results
        
    def score(item):
        s = 0
        title = str(item.get("title", "")).lower()
        # authors might be list or string depending on where it came from
        authors_val = item.get("authors", [])
        if isinstance(authors_val, list):
            authors = " ".join(authors_val).lower()
        else:
            authors = str(authors_val).lower()
            
        q = query.lower()
        
        # Exact title match
        if title == q:
            s += 100
        # Query in title
        elif q in title:
            s += 50
        
        # Recent publications (bonus)
        try:
            year_str = str(item.get("year", "0"))
            # Handle "2023a" or similar
            year = int("".join(filter(str.isdigit, year_str)) or 0)
            if year > 2020:
                s += 5
        except:
            pass
            
        # Has DOI (quality signal)
        if item.get("doi"):
            s += 2
            
        return s
        
    return sorted(results, key=score, reverse=True)

########################################
# LOOKUP WRAPPERS
########################################

def lookup_works(query_text: str, cache: Dict = None, limit: int = 5, **kwargs) -> List[Dict[str, Any]]:
    """
    Look up multiple works (wrapper around ReferenceManager).
    """
    if not query_text or not query_text.strip():
        logger.warning("Empty query text provided")
        return []
        
    manager = _get_manager()
    
    # Sync cache: update manager with provided cache
    if cache:
        manager.cache.update(cache)
        
    # Perform search
    results = manager.search_works(query_text, limit, **kwargs)
    
    # Sync cache back
    if cache is not None:
        cache.update(manager.cache)
    
    return [asdict(r) if is_dataclass(r) else r for r in results]

def lookup_single_work(query_text: str, cache: Dict = None) -> Optional[Dict[str, Any]]:
    """
    Look up a single work (wrapper around ReferenceManager).
    """
    results = lookup_works(query_text, cache, limit=1)
    return results[0] if results else None

def lookup_author_works(author_name: str, cache: Dict = None, **kwargs) -> List[Dict[str, Any]]:
    """
    Look up works by author (wrapper around ReferenceManager).
    """
    if not author_name or not author_name.strip():
        logger.warning("Empty author name provided")
        return []
        
    manager = _get_manager()
    
    # Sync cache
    if cache:
        manager.cache.update(cache)
        
    # Perform search
    results = manager.search_author_works(author_name, **kwargs)
    
    # Sync cache back
    if cache is not None:
        cache.update(manager.cache)
    
    # Convert to dicts and rank
    results_dict = [asdict(r) if is_dataclass(r) else r for r in results]
    
    return rank_results(results_dict, author_name)
########################################
# FORMATTING HELPERS PER STYLE
########################################

def _authors_as_string_harvard(authors):
    # "Smith, J. and Doe, A." / "Smith, J., Doe, A. and Lee, K."
    if not authors:
        return "UNKNOWN"
    if len(authors) == 1:
        return authors[0]
    if len(authors) == 2:
        return " and ".join(authors)
    return ", ".join(authors[:-1]) + " and " + authors[-1]

def _authors_as_string_apa(authors):
    # "Smith, J., & Doe, A." / "Smith, J., Doe, A., & Lee, K."
    if not authors:
        return "UNKNOWN"
    if len(authors) == 1:
        return authors[0]
    if len(authors) == 2:
        return authors[0] + ", & " + authors[1]
    # Oxford comma style with &
    return ", ".join(authors[:-1]) + ", & " + authors[-1]

def _authors_as_string_ieee(authors):
    # "J. Smith and A. Doe" / "J. Smith, A. Doe, and K. Lee"
    # Need to flip "Surname, Given" into "G. Surname"
    if not authors:
        return "UNKNOWN"
    def flip(a):
        # "Ruvinga, Stenford" -> "S. Ruvinga"
        if "," in a:
            fam, giv = [x.strip() for x in a.split(",", 1)]
        else:
            # fallback
            parts = a.split()
            if len(parts) == 1:
                return parts[0]
            fam = parts[-1]
            giv = " ".join(parts[:-1])
        initials = " ".join([p[0] + "." for p in giv.split() if p])
        return f"{initials} {fam}".strip()

    flipped = [flip(x) for x in authors]

    if len(flipped) == 1:
        return flipped[0]
    if len(flipped) == 2:
        return flipped[0] + " and " + flipped[1]
    return ", ".join(flipped[:-1]) + ", and " + flipped[-1]

def _authors_as_string_mla(authors):
    # MLA: "Surname, First Name" for first author, "First Name Surname" for others
    # "Smith, John, and Jane Doe" / "Smith, John, Jane Doe, and Bob Lee"
    if not authors:
        return "UNKNOWN"
    
    def format_first(a):
        # Keep as "Surname, Given"
        return a
    
    def format_rest(a):
        # "Surname, Given" -> "Given Surname"
        if "," in a:
            fam, giv = [x.strip() for x in a.split(",", 1)]
            return f"{giv} {fam}"
        return a
    
    if len(authors) == 1:
        return format_first(authors[0])
    
    formatted = [format_first(authors[0])] + [format_rest(a) for a in authors[1:]]
    
    if len(formatted) == 2:
        return formatted[0] + ", and " + formatted[1]
    return ", ".join(formatted[:-1]) + ", and " + formatted[-1]

def _authors_as_string_chicago(authors):
    # Chicago: "Surname, First Name, and First Name Surname"
    # Same as MLA essentially
    if not authors:
        return "UNKNOWN"
    
    def format_first(a):
        return a
    
    def format_rest(a):
        if "," in a:
            fam, giv = [x.strip() for x in a.split(",", 1)]
            return f"{giv} {fam}"
        return a
    
    if len(authors) == 1:
        return format_first(authors[0])
    
    formatted = [format_first(authors[0])] + [format_rest(a) for a in authors[1:]]
    
    if len(formatted) == 2:
        return formatted[0] + " and " + formatted[1]
    return ", ".join(formatted[:-1]) + ", and " + formatted[-1]

def _authors_as_string_vancouver(authors):
    # Vancouver: "Surname Initial, Surname Initial"
    # "Smith J, Doe A, Lee K"
    if not authors:
        return "UNKNOWN"
    
    def format_vancouver(a):
        # "Ruvinga, Stenford" -> "Ruvinga S"
        if "," in a:
            fam, giv = [x.strip() for x in a.split(",", 1)]
        else:
            parts = a.split()
            if len(parts) == 1:
                return parts[0]
            fam = parts[-1]
            giv = " ".join(parts[:-1])
        initials = "".join([p[0].upper() for p in giv.split() if p])
        return f"{fam} {initials}".strip()
    
    formatted = [format_vancouver(a) for a in authors]
    return ", ".join(formatted)

########################################
# IN-TEXT CITATION PER STYLE
########################################

def in_text_citation(meta, style, index_number=None):
    """
    style: 'harvard', 'apa', 'ieee', 'mla', 'chicago', 'vancouver'
    index_number: only used for IEEE and Vancouver, e.g. [3]
    """
    authors = meta["authors"]
    year = meta["year"]

    if style == "ieee" or style == "vancouver":
        # IEEE and Vancouver use numeric in-text: [1]
        return f"[{index_number if index_number is not None else '?'}]"

    if not authors:
        return "(UNKNOWN)"

    first_author_surname = (authors[0].split(",")[0]).strip()

    if style == "mla":
        # MLA: (Surname) or (Surname and Surname) or (Surname et al.)
        if len(authors) > 2:
            return f"({first_author_surname} et al.)"
        elif len(authors) == 2:
            second_surname = (authors[1].split(",")[0]).strip()
            return f"({first_author_surname} and {second_surname})"
        else:
            return f"({first_author_surname})"

    # Harvard / APA / Chicago all use 'Surname (year)' style
    if len(authors) > 2:
        # Harvard: (Surname et al., 2022)
        # APA:     (Surname et al., 2022)
        # Chicago: (Surname et al. 2022)
        if style == "chicago":
            return f"({first_author_surname} et al. {year})"
        else:
            return f"({first_author_surname} et al., {year})"

    if len(authors) == 2:
        second_surname = (authors[1].split(",")[0]).strip()
        if style == "apa":
            return f"({first_author_surname} & {second_surname}, {year})"
        elif style == "chicago":
            return f"({first_author_surname} and {second_surname} {year})"
        else:
            return f"({first_author_surname} and {second_surname}, {year})"

    # single author
    if style == "chicago":
        return f"({first_author_surname} {year})"
    else:
        return f"({first_author_surname}, {year})"

########################################
# FULL REFERENCE BUILDER PER STYLE
########################################

def reference_entry(meta, style, index_number=None):
    """
    Build the full bibliography entry for one item,
    given the chosen style.
    """
    pub_type = meta.get("pub_type", "").lower()
    is_article_like = (
        pub_type in ["journal-article", "article-journal", "proceedings-article"]
        or bool(meta.get("journal"))
    )

    if style == "harvard":
        if is_article_like:
            # Author (Year) 'Title', Journal, vol(issue), pp. pages. doi:...
            author_str = _authors_as_string_harvard(meta["authors"])
            vol_issue = ""
            if meta["volume"] and meta["issue"]:
                vol_issue = f"{meta['volume']}({meta['issue']})"
            elif meta["volume"]:
                vol_issue = meta["volume"]
            elif meta["issue"]:
                vol_issue = f"({meta['issue']})"
            pages = f"pp. {meta['pages']}" if meta["pages"] else ""
            out = f"{author_str} ({meta['year']}) '{meta['title']}', {meta['journal']}"
            if vol_issue:
                out += f", {vol_issue}"
            if pages:
                out += f", {pages}"
            if meta["doi"]:
                out += f". doi:{meta['doi']}"
            else:
                out += "."
            return out
        else:
            # books
            author_str = _authors_as_string_harvard(meta.get("authors", []))
            pub_block = meta.get("publisher", "")
            out = f"{author_str} ({meta.get('year', 'n.d.')}) {meta.get('title', 'Untitled')}."
            
            if pub_type in ["web", "web-page"]:
                # Surname, Initial. (Year) Title. Available at: URL (Accessed: date).
                out = f"{author_str} ({meta.get('year', 'n.d.')}) {meta.get('title', 'Untitled')}."
                if meta.get("url"):
                    out += f" Available at: {meta['url']}"
                    if meta.get("access_date"):
                        out += f" (Accessed: {meta['access_date']})"
                    out += "."
                return out

            if pub_block:
                out += f" {pub_block}."
            return out

    if style == "apa":
        if is_article_like:
            # Author. (Year). Title. Journal, vol(issue), pages. https://doi.org/...
            author_str = _authors_as_string_apa(meta["authors"])
            vol_issue = ""
            if meta["volume"] and meta["issue"]:
                vol_issue = f"{meta['volume']}({meta['issue']})"
            elif meta["volume"]:
                vol_issue = meta["volume"]
            elif meta["issue"]:
                vol_issue = f"({meta['issue']})"
            pages = meta["pages"] if meta["pages"] else ""
            out = f"{author_str} ({meta['year']}). {meta['title']}. {meta['journal']}"
            if vol_issue:
                out += f", {vol_issue}"
            if pages:
                out += f", {pages}"
            if meta["doi"]:
                out += f". https://doi.org/{meta['doi']}"
            else:
                out += "."
            return out
        else:
            # books
            author_str = _authors_as_string_apa(meta.get("authors", []))
            out = f"{author_str} ({meta.get('year', 'n.d.')}). {meta.get('title', 'Untitled')}."
            if meta.get("publisher"):
                out += f" {meta['publisher']}."
            return out

    if style == "ieee":
        # IEEE: [n] J. Smith and A. Doe, "Title," Journal, vol. 10, no. 2, pp. 33–41, 2022.
        # Books are similar but without journal/vol/no/pp.
        num_label = f"[{index_number if index_number is not None else '?'}]"
        author_str = _authors_as_string_ieee(meta["authors"])
        year = meta["year"]

        if is_article_like:
            bits = [f"{num_label} {author_str}, \"{meta['title']}\","]
            if meta["journal"]:
                bits.append(meta["journal"] + ",")
            if meta["volume"]:
                bits.append(f"vol. {meta['volume']},")
            if meta["issue"]:
                bits.append(f"no. {meta['issue']},")
            if meta["pages"]:
                bits.append(f"pp. {meta['pages']},")
            bits.append(year + ".")
            return " ".join(bits)
        else:
            # Book in IEEE:
            # [n] J. Smith and A. Doe, Title, Publisher, 2022.
            pub_name = meta.get("publisher", "")
            bits = [f"{num_label} {author_str}, {meta.get('title', 'Untitled')}, {pub_name}, {year}."]
            return " ".join(bits)

    if style == "mla":
        # MLA 9th edition: Surname, First Name, and First Name Surname. "Title." Journal, vol. X, no. Y, Year, pp. Z-ZZ.
        # Books: Surname, First Name. Title. Publisher, Year.
        author_str = _authors_as_string_mla(meta["authors"])
        year = meta["year"]
        
        if is_article_like:
            out = f"{author_str}. \"{meta['title']}\". {meta['journal']}"
            if meta["volume"]:
                out += f", vol. {meta['volume']}"
            if meta["issue"]:
                out += f", no. {meta['issue']}"
            out += f", {year}"
            if meta["pages"]:
                out += f", pp. {meta['pages']}"
            out += "."
            return out
        else:
            # Book
            out = f"{author_str}. {meta.get('title', 'Untitled')}."
            if meta.get("publisher"):
                out += f" {meta['publisher']}, {year}."
            else:
                out += f" {year}."
            return out

    if style == "chicago":
        # Chicago (Author-Date): Surname, First Name, and First Name Surname. Year. "Title." Journal vol (issue): pages.
        # Books: Surname, First Name. Year. Title. Publisher.
        author_str = _authors_as_string_chicago(meta["authors"])
        year = meta["year"]
        
        if is_article_like:
            out = f"{author_str}. {year}. \"{meta['title']}\". {meta['journal']}"
            if meta["volume"]:
                out += f" {meta['volume']}"
            if meta["issue"]:
                out += f" ({meta['issue']})"
            if meta["pages"]:
                out += f": {meta['pages']}"
            out += "."
            if meta["doi"]:
                out += f" https://doi.org/{meta['doi']}"
            return out
        else:
            # Book
            out = f"{author_str}. {year}. {meta.get('title', 'Untitled')}."
            if meta.get("publisher"):
                out += f" {meta['publisher']}."
            return out

    if style == "vancouver":
        # Vancouver: [n] Surname Initial, Surname Initial. Title. Journal. Year;vol(issue):pages.
        # Books: [n] Surname Initial, Surname Initial. Title. Publisher; Year.
        num_label = f"{index_number if index_number is not None else '?'}."
        author_str = _authors_as_string_vancouver(meta["authors"])
        year = meta["year"]
        
        if is_article_like:
            out = f"{num_label} {author_str}. {meta['title']}. {meta['journal']}. {year}"
            if meta["volume"]:
                out += f";{meta['volume']}"
            if meta["issue"]:
                out += f"({meta['issue']})"
            if meta["pages"]:
                out += f":{meta['pages']}"
            out += "."
            return out
        else:
            # Book
            out = f"{num_label} {author_str}. {meta.get('title', 'Untitled')}."
            if meta.get("publisher"):
                out += f" {meta['publisher']}; {year}."
            else:
                out += f" {year}."
            return out

    # fallback
    return "UNKNOWN STYLE"

@handle_api_errors(max_retries=3, backoff_factor=0.5)
def search_semantic_scholar_single(query_text: str, limit: int = 5) -> Optional[Dict[str, Any]]:
    """
    Search for a single work using Semantic Scholar API (free, no key required).
    
    Args:
        query_text: The search query
        limit: Maximum number of results to return (default: 5)
        
    Returns:
        Dictionary containing normalized metadata or None if an error occurs
    """
    url = "https://api.semanticscholar.org/graph/v1/paper/search"
    params = {
        'query': query_text,
        'limit': limit,
        'fields': 'title,authors,year,venue,publicationTypes,externalIds,publicationDate'
    }
    
    try:
        response = requests.get(url, params=params, timeout=10)
        response.raise_for_status()
        data = response.json()
        
        if not data.get('data') or len(data['data']) == 0:
            logger.info(f"No results from Semantic Scholar for: {query_text}")
            return None
            
        # Take the first result
        paper = data['data'][0]
        
        # Extract authors
        authors = []
        for author in paper.get('authors', []):
            name = author.get('name', '')
            if name:
                # Convert "First Last" to "Last, First"
                parts = name.split()
                if len(parts) >= 2:
                    surname = parts[-1]
                    given = " ".join(parts[:-1])
                    authors.append(f"{surname}, {given}")
                else:
                    authors.append(name)
        
        # Extract year
        year = paper.get('year', 'n.d.')
        if not year:
            pub_date = paper.get('publicationDate', '')
            year = pub_date.split('-')[0] if pub_date else 'n.d.'
        
        # Extract DOI
        doi = ""
        external_ids = paper.get('externalIds', {})
        if external_ids:
            doi = external_ids.get('DOI', '')
        
        return {
            "source": "semantic_scholar",
            "pub_type": "journal-article",
            "authors": authors,
            "year": str(year),
            "title": paper.get('title', 'NO TITLE'),
            "journal": paper.get('venue', ''),
            "publisher": "",
            "location": "",
            "volume": "",
            "issue": "",
            "pages": "",
            "doi": doi,
        }
        
    except Exception as e:
        logger.error(f"Error searching Semantic Scholar: {str(e)}")
        return None

@handle_api_errors(max_retries=3, backoff_factor=0.5)
def search_college_catalog(query_text: str, max_results: int = 5) -> Optional[List[Dict[str, Any]]]:
    """
    Search the college catalog for works.
    
    Args:
        query_text: The search query
        max_results: Maximum number of results to return (default: 5)
        
    Returns:
        List of dictionaries containing work metadata or None if an error occurs
    """
    try:
        params = {
            'q': query_text,
            'per_page': max_results
        }
        
        response = requests.get(
            f"{COLLEGE_CATALOG_API_URL}/search",
            params=params,
            timeout=10
        )
        response.raise_for_status()
        
        data = response.json()
        results = []
        
        for item in data.get('items', [])[:max_results]:
            # Map college catalog fields to our standard format
            authors = item.get('authors', [])
            if not isinstance(authors, list):
                authors = [str(authors)] if authors else []
                
            result = {
                'title': item.get('title', 'Untitled'),
                'authors': authors,
                'year': item.get('publication_year'),
                'journal': item.get('publisher', ''),
                'volume': item.get('volume', ''),
                'issue': item.get('issue', ''),
                'pages': str(item.get('pages', '')),
                'doi': item.get('doi', ''),
                'url': f"{COLLEGE_CATALOG_API_URL}/items/{item.get('id', '')}",
                'pub_type': item.get('item_type', 'Book').lower(),
                'source': 'college_catalog',
                'abstract': item.get('description', '')
            }
            results.append(result)
            
        return results if results else None
        
    except Exception as e:
        logger.warning(f"Error searching college catalog: {str(e)}")
        return None

def search_arxiv_single(query_text: str, max_results: int = 5) -> Optional[Dict[str, Any]]:
    """
    Search for a single work using arXiv API (free, no key required).
    
    Args:
        query_text: The search query
        max_results: Maximum number of results to return (default: 5)
        
    Returns:
        Dictionary containing normalized metadata or None if an error occurs
    """
    import xml.etree.ElementTree as ET
    
    url = "http://export.arxiv.org/api/query"
    params = {
        'search_query': f'all:{query_text}',
        'start': 0,
        'max_results': max_results
    }
    
    try:
        response = requests.get(url, params=params, timeout=10)
        response.raise_for_status()
        
        # Parse XML response
        root = ET.fromstring(response.content)
        
        # arXiv uses Atom namespace
        ns = {'atom': 'http://www.w3.org/2005/Atom'}
        
        entries = root.findall('atom:entry', ns)
        if not entries:
            logger.info(f"No results from arXiv for: {query_text}")
            return None
        
        # Take the first result
        entry = entries[0]
        
        # Extract title
        title_elem = entry.find('atom:title', ns)
        title = title_elem.text.strip() if title_elem is not None else 'NO TITLE'
        
        # Extract authors
        authors = []
        for author in entry.findall('atom:author', ns):
            name_elem = author.find('atom:name', ns)
            if name_elem is not None:
                name = name_elem.text.strip()
                # Convert "First Last" to "Last, First"
                parts = name.split()
                if len(parts) >= 2:
                    surname = parts[-1]
                    given = " ".join(parts[:-1])
                    authors.append(f"{surname}, {given}")
                else:
                    authors.append(name)
        
        # Extract publication date
        published_elem = entry.find('atom:published', ns)
        year = 'n.d.'
        if published_elem is not None:
            pub_date = published_elem.text.strip()
            year = pub_date.split('-')[0] if '-' in pub_date else pub_date[:4]
        
        # Extract arXiv ID for DOI-like reference
        id_elem = entry.find('atom:id', ns)
        arxiv_id = ""
        if id_elem is not None:
            arxiv_url = id_elem.text.strip()
            arxiv_id = arxiv_url.split('/')[-1]  # Extract ID from URL
        
        return {
            "source": "arxiv",
            "pub_type": "preprint",
            "authors": authors,
            "year": year,
            "title": title,
            "journal": "arXiv preprint",
            "publisher": "",
            "location": "",
            "volume": "",
            "issue": "",
            "pages": "",
            "doi": f"arXiv:{arxiv_id}" if arxiv_id else "",
        }
        
    except Exception as e:
        logger.error(f"Error searching arXiv: {str(e)}")
        return None

########################################
# EXPORT FUNCTIONS
########################################

def export_bibtex(refs: List[Dict[str, Any]]) -> str:
    """
    Export references in BibTeX format.
    
    Args:
        refs: List of reference dictionaries
        
    Returns:
        str: References formatted in BibTeX
    """
    bibtex_entries = []
    
    for i, ref in enumerate(refs):
        # Generate a unique key for the entry
        first_author = ''
        if ref.get('authors'):
            if isinstance(ref['authors'], str):
                first_author = ref['authors'].split(',')[0].strip()
            elif isinstance(ref['authors'], list) and len(ref['authors']) > 0:
                first_author = ref['authors'][0].split(',')[0].strip()
        
        year = ref.get('year', 'n.d.')
        entry_key = f"{first_author}{year}".lower().replace(' ', '')
        
        # Determine entry type
        entry_type = 'article'  # default
        if ref.get('pub_type'):
            if 'book' in ref['pub_type'].lower():
                entry_type = 'book'
            elif 'thesis' in ref['pub_type'].lower():
                entry_type = 'phdthesis'
        
        # Start entry
        entry = [f"@{entry_type}{{{entry_key},"]
        
        # Add fields
        if 'title' in ref:
            entry.append(f'    title = {{{ref["title"]}}},')
            
        if 'authors' in ref:
            if isinstance(ref['authors'], list):
                authors = ' and '.join(ref['authors'])
            else:
                authors = ref['authors']
            entry.append(f'    author = {{{authors}}},')
            
        if 'year' in ref:
            entry.append(f'    year = {{{ref["year"]}}},')
            
        if 'journal' in ref:
            entry.append(f'    journal = {{{ref["journal"]}}},')
            
        if 'volume' in ref:
            entry.append(f'    volume = {{{ref["volume"]}}},')
            
        if 'issue' in ref:
            entry.append(f'    number = {{{ref["issue"]}}},')
            
        if 'pages' in ref:
            entry.append(f'    pages = {{{ref["pages"]}}},')
            
        if 'doi' in ref:
            entry.append(f'    doi = {{{ref["doi"]}}},')
            
        if 'url' in ref:
            entry.append(f'    url = {{{ref["url"]}}},')
        elif 'doi' in ref:
            entry.append(f'    url = {{https://doi.org/{ref["doi"]}}},')
            
        if 'publisher' in ref:
            entry.append(f'    publisher = {{{ref["publisher"]}}},')
            
        # Remove trailing comma from last field
        if entry[-1].endswith(','):
            entry[-1] = entry[-1][:-1]
            
        entry.append('}')
        bibtex_entries.append('\n'.join(entry))
    
    return '\n\n'.join(bibtex_entries)

def export_ris(refs: List[Dict[str, Any]]) -> str:
    """
    Export references in RIS format.
    
    Args:
        refs: List of reference dictionaries
        
    Returns:
        str: References formatted in RIS
    """
    ris_entries = []
    
    for ref in refs:
        entry = []
        
        # Determine reference type
        ref_type = 'JOUR'  # default to journal article
        if ref.get('pub_type'):
            if 'book' in ref['pub_type'].lower():
                ref_type = 'BOOK'
            elif 'thesis' in ref['pub_type'].lower():
                ref_type = 'THES'
        entry.append(f'TY  - {ref_type}')
        
        # Add fields
        if 'title' in ref:
            entry.append(f'TI  - {ref["title"]}')
            
        if 'authors' in ref:
            if isinstance(ref['authors'], list):
                for author in ref['authors']:
                    entry.append(f'AU  - {author}')
            else:
                entry.append(f'AU  - {ref["authors"]}')
                
        if 'year' in ref:
            entry.append(f'PY  - {ref["year"]}')
            
        if 'journal' in ref:
            entry.append(f'JO  - {ref["journal"]}')
            
        if 'volume' in ref:
            entry.append(f'VL  - {ref["volume"]}')
            
        if 'issue' in ref:
            entry.append(f'IS  - {ref["issue"]}')
            
        if 'pages' in ref:
            entry.append(f'SP  - {ref["pages"].split("-")[0]}')
            entry.append(f'EP  - {ref["pages"].split("-")[-1]}')
            
        if 'doi' in ref:
            entry.append(f'DO  - {ref["doi"]}')
            
        if 'url' in ref:
            entry.append(f'UR  - {ref["url"]}')
        elif 'doi' in ref:
            entry.append(f'UR  - https://doi.org/{ref["doi"]}')
            
        if 'publisher' in ref:
            entry.append(f'PB  - {ref["publisher"]}')
            
        # End of reference
        entry.append('ER  - ')
        entry.append('')
        
        ris_entries.append('\n'.join(entry))
    
    return '\n'.join(ris_entries)

########################################
# DEDUPE/SORT
########################################

def get_dedupe_key(m):
    """
    Generates a deterministic key for a reference.
    Uses DOI if present, otherwise a combination of title, first author surname, and year.
    Supports both dicts and objects.
    
    Collision Risk: 
    - Different DOIs for the same work (rare) will bypass this.
    - Identical titles/authors/years for different works (e.g. 'Letter to Editor') might clash.
    """
    def _get(obj, k):
        if isinstance(obj, dict):
            return obj.get(k)
        return getattr(obj, k, None)

    doi = _get(m, "doi")
    if doi:
        return doi.lower().strip()
    
    # Normalize title: lowercase, alphanumeric only, trimmed
    title_val = _get(m, "title")
    title = "".join(c for c in (title_val or "").lower() if c.isalnum())
    
    # First author surname
    author = ""
    authors = _get(m, "authors")
    if authors:
        if isinstance(authors, list) and len(authors) > 0:
            author = authors[0].split(",")[0].lower().strip()
        elif isinstance(authors, str):
            author = authors.split(",")[0].lower().strip()
            
    year = str(_get(m, "year") or "").strip()
    
    return f"{title}|{author}|{year}"

def validate_publication(data):
    """
    Validate that a dictionary represents a valid publication object.
    
    Checks:
    - Input must be a dictionary.
    - 'title' is required and must not be empty.
    - 'authors' must be a list or string if present.
    - 'year' must be convertible to string/int if present.
    
    Returns:
        (is_valid, error_message)
    """
    if not isinstance(data, dict):
        return False, "Publication data must be a JSON object (dictionary)."
        
    # Required fields
    if not data.get("title") or not str(data["title"]).strip():
        return False, "Publication must have a 'title'."
        
    # Type checks
    if "authors" in data:
        if not isinstance(data["authors"], (list, str)):
            return False, "'authors' must be a list of strings or a single string."
            
    if "year" in data and data["year"]:
        # Just check if it's broadly sensible (string/int)
        if not isinstance(data["year"], (str, int)):
            return False, "'year' must be a string or number."
            
    return True, None

def dedupe(collected):
    """Keep first occurrence of each item using the deterministic key strategy."""
    seen = set()
    uniq = []
    for m in collected:
        key = get_dedupe_key(m)
        if key not in seen:
            uniq.append(m)
            seen.add(key)
    return uniq

def is_duplicate(new_ref, existing_refs):
    """Check if a new reference is already present in a collection of existing references."""
    new_key = get_dedupe_key(new_ref)
    for ref in existing_refs:
        if get_dedupe_key(ref) == new_key:
            return True
    return False

def sort_for_bibliography(uniq, style):
    """Harvard/APA/MLA/Chicago: alphabetical. IEEE/Vancouver: in added order."""
    if style == "ieee" or style == "vancouver":
        return uniq  # keep original order of addition (numbered styles)
    # alphabetical by first author's surname
    def sort_key(x):
        if x["authors"]:
            return x["authors"][0].split(",")[0].lower()
        return x["title"].lower()
    return sorted(uniq, key=sort_key)

def save_references_to_word(refs_sorted, folder, filename, sources, style):
    os.makedirs(folder, exist_ok=True)
    doc = Document()
    doc.add_heading("References", level=1)
    doc.add_paragraph(
        f"Generated {datetime.now():%Y-%m-%d %H:%M} "
        f"using: {', '.join(sorted(list(sources)))}. "
        f"Style: {style.upper()}. "
        "Some fields (publisher, DOI, year) may require manual checking."
    )
    doc.add_paragraph("")
    for idx, m in enumerate(refs_sorted, start=1):
        entry_text = reference_entry(
            m,
            style=style,
            index_number=idx if (style == "ieee" or style == "vancouver") else None
        )
        doc.add_paragraph(entry_text)
    path = os.path.join(folder, filename)
    doc.save(path)
    return path

########################################
# MENU ACTIONS
########################################

def action_search_single(refs, cache, sources, style):
    q = input("Enter title or keywords: ").strip()
    if not q:
        return
    m = lookup_single_work(q, cache)
    if not m:
        print("No match found.")
        return
    sources.add(m["source"])

    # show in-text + ref in chosen style
    print("\nIn-text:", in_text_citation(m, style, index_number=1))
    print("Full ref:", reference_entry(m, style=style, index_number=1))

    add_q = input("Add this to bibliography? (y/n): ").strip().lower()
    if add_q.startswith("y"):
        refs.append(m)
        print("Added.")

def action_search_author(refs, cache, sources, style):
    a = input("Enter author name: ").strip()
    if not a:
        return
    works = lookup_author_works(a, cache)
    if not works:
        print("No works found.")
        return
    for w in works:
        sources.add(w["source"])

    while True:
        print(f"\nWorks for {a}:")
        for i, m in enumerate(works, 1):
            j = m["journal"] or m["publisher"]
            print(f"{i}. {m['title']} ({m['year']}) — {j}")
        s = input("Select numbers (comma) or Enter to finish: ").strip()
        if not s:
            break
        picks = [x.strip() for x in s.split(",") if x.strip().isdigit()]
        for p in picks:
            idx = int(p)
            if 1 <= idx <= len(works):
                chosen = works[idx-1]
                refs.append(chosen)
                print("Added:")
                print("  In-text:", in_text_citation(chosen, style, index_number=1))
                print("  Ref   :", reference_entry(chosen, style=style, index_number=1))
            else:
                print(f"Ignored {p}: out of range.")

def action_view_current(refs, style):
    uniq = dedupe(refs)
    if not uniq:
        print("\n(No items yet.)")
        return
    refs_sorted = sort_for_bibliography(uniq, style)
    print("\nCurrent bibliography:")
    for idx, m in enumerate(refs_sorted, start=1):
        print(f"{idx}. {reference_entry(m, style=style, index_number=idx)}")

def action_view_citations(refs, style):
    uniq = dedupe(refs)
    if not uniq:
        print("\n(No items yet.)")
        return
    refs_sorted = sort_for_bibliography(uniq, style)
    print("\nCitations and references:")
    for idx, m in enumerate(refs_sorted, start=1):
        print(f"[{idx}] In-text: {in_text_citation(m, style, index_number=idx)}")
        print(f"     Ref   : {reference_entry(m, style=style, index_number=idx)}")
        print("")

def action_export_and_exit(refs, sources, style):
    uniq = dedupe(refs)
    if not uniq:
        print("Nothing to save.")
        return True
    refs_sorted = sort_for_bibliography(uniq, style)
    path = save_references_to_word(refs_sorted, DOWNLOAD_FOLDER, WORD_FILENAME, sources, style)
    print("Saved to:", path)
    return True

def action_change_style(current_style):
    print("\nChoose style:")
    print("1. Harvard")
    print("2. APA")
    print("3. IEEE")
    print("4. MLA")
    print("5. Chicago")
    print("6. Vancouver")
    choice = input("Style (1-6): ").strip()
    if choice == "1":
        return "harvard"
    if choice == "2":
        return "apa"
    if choice == "3":
        return "ieee"
    if choice == "4":
        return "mla"
    if choice == "5":
        return "chicago"
    if choice == "6":
        return "vancouver"
    print("Keeping previous style.")
    return current_style

########################################
# MAIN APP LOOP
########################################

def main():
    cache = load_cache()
    refs = []
    sources = set()
    style = "harvard"  # default style

    print("===== Reference Assistant =====")
    print(f"Saving to: {DOWNLOAD_FOLDER}/{WORD_FILENAME}")
    print(f"Cache file: {CACHE_FILE}")
    print(f"Initial style: {style.upper()}")

    while True:
        print("\nMenu:")
        print(f"1. Search by title/book (current style: {style.upper()})")
        print(f"2. Search by author (current style: {style.upper()})")
        print("3. View current bibliography")
        print("4. Export & exit")
        print("5. Exit without saving")
        print("6. Change reference style")
        print("7. Show citations for current list")

        c = input("Choice (1-7): ").strip()

        if c == "1":
            action_search_single(refs, cache, sources, style)
        elif c == "2":
            action_search_author(refs, cache, sources, style)
        elif c == "3":
            action_view_current(refs, style)
        elif c == "4":
            should_quit = action_export_and_exit(refs, sources, style)
            if should_quit:
                break
        elif c == "5":
            print("Exiting without saving.")
            break
        elif c == "6":
            style = action_change_style(style)
        elif c == "7":
            action_view_citations(refs, style)
        else:
            print("Invalid choice.")

if __name__ == "__main__":
    main()
