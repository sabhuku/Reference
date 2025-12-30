from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file, jsonify, Response
import json
import os
import io
import sys
import csv
from functools import wraps
from datetime import datetime
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from io import StringIO

# Import the reference model
from src.referencing.models import Reference, Author

# Ensure project root (one level up) is on sys.path so top-level modules like
# `referencing.py` can be imported when running this script from the `ui/`
# folder.
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
if ROOT not in sys.path:
    sys.path.insert(0, ROOT)

from src.referencing import referencing

# Persistence for manual refs
PERSIST_PATH = os.path.join(ROOT, 'manual_refs.json')

def load_persisted_refs():
    try:
        if os.path.exists(PERSIST_PATH):
            with open(PERSIST_PATH, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if isinstance(data, list):
                    return data
    except Exception:
        pass
    return []

def save_persisted_refs(refs):
    try:
        with open(PERSIST_PATH, 'w', encoding='utf-8') as f:
            json.dump(refs, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

import re

def normalize_authors(raw: str):
    """Normalize a free-form authors string into a list of 'Surname, Given' entries.

    Heuristics supported:
    - Semicolon-separated or newline separated lists
    - 'and' separated lists
    - Comma-paired tokens like 'Smith, John, Doe, Jane' -> ['Smith, John', 'Doe, Jane']
    - Space-separated names 'John Smith, Mary Jones' -> ['Smith, John', 'Jones, Mary']
    - Single author in any form
    """
    if not raw:
        return []
    s = raw.strip()
    # Split into candidate author strings
    if ';' in s:
        parts = [p.strip() for p in s.split(';') if p.strip()]
    elif '\n' in s:
        parts = [p.strip() for p in s.splitlines() if p.strip()]
    elif re.search(r'\band\b', s, flags=re.I):
        parts = [p.strip() for p in re.split(r'\band\b', s, flags=re.I) if p.strip()]
    else:
        tokens = [t.strip() for t in s.split(',') if t.strip()]
        # if tokens look like pairs (surname, given, surname, given)
        # only treat as pairs when the surname tokens (even indices) look like single-token surnames
        if len(tokens) >= 2 and len(tokens) % 2 == 0:
            even_tokens_are_single = all(' ' not in tokens[i] for i in range(0, len(tokens), 2))
            if even_tokens_are_single:
                parts = []
                for i in range(0, len(tokens), 2):
                    parts.append(tokens[i] + ', ' + tokens[i+1])
            else:
                # special-case: 'Surname Suffix, Given' should be treated as a paired entry
                combined = False
                if len(tokens) == 2:
                    last_token_of_first = tokens[0].split()[-1]
                    if last_token_of_first.rstrip('.') in {s.rstrip('.') for s in ("Jr", "Jr.", "Sr", "Sr.") }:
                        parts = [tokens[0] + ', ' + tokens[1]]
                        combined = True
                if not combined:
                    # If tokens look like full names (contain spaces), treat each as an author
                    if len(tokens) > 1 and all(' ' in t for t in tokens):
                        parts = tokens
                    else:
                        parts = [p.strip() for p in s.split(',') if p.strip()]
        else:
            # If there are multiple comma-separated tokens but they contain spaces, assume 'Given Surname' tokens
            if len(tokens) > 1 and all(' ' in t for t in tokens):
                parts = tokens
            else:
                # fallback: treat original string as single or multiple authors separated by comma
                parts = [p.strip() for p in s.split(',') if p.strip()]

    # common surname particles to help detect multi-token family names
    particles = {"van", "von", "de", "del", "da", "di", "la", "le", "du", "st", "st.", "der"}
    suffixes = {"Jr", "Jr.", "Sr", "Sr.", "II", "III", "IV"}

    def fmt(a: str) -> str:
        if ',' in a:
            family, given = [x.strip() for x in a.split(',', 1)]
            return f"{family}, {given}"
        parts = a.split()
        if len(parts) == 1:
            return parts[0]

        # If the last token looks like an initial (e.g. 'J.' or 'J'), assume
        # the form is 'Family Initial' and treat the first token as family.
        last = parts[-1]
        last_stripped = last.rstrip('.')
        # don't treat common suffixes as initials
        suffix_stripped = {s.rstrip('.') for s in suffixes}
        is_initial_like = (
            (last.endswith('.') or len(last_stripped) <= 2)
            and last_stripped.isalpha()
            and last_stripped not in suffix_stripped
        )

        if is_initial_like and len(parts) >= 2:
            family = parts[0]
            given = ' '.join(parts[1:])
        else:
            # handle suffixes at the end like 'John Smith Jr.' -> family 'Smith Jr.'
            if parts[-1] in suffixes and len(parts) >= 3:
                family = parts[-2] + ' ' + parts[-1]
                given = ' '.join(parts[:-2])
            else:
                # try to detect particles that are part of the family name, e.g. 'van der Waals'
                if len(parts) >= 3:
                    # build family tokens by scanning backwards for known particles
                    i = len(parts) - 2
                    family_tokens = [parts[-1]]
                    while i >= 0 and parts[i].lower() in particles:
                        family_tokens.insert(0, parts[i])
                        i -= 1
                    if len(family_tokens) > 1:
                        family = ' '.join(family_tokens)
                        given = ' '.join(parts[:i+1]) if i >= 0 else ''
                    else:
                        # default: treat final token as family name
                        family = parts[-1]
                        given = ' '.join(parts[:-1])
                else:
                    family = parts[-1]
                    given = ' '.join(parts[:-1])

        return f"{family}, {given}"

    return [fmt(p) for p in parts]

app = Flask(__name__)

# Configuration
app.config.update(
    # For local use only — change in production
    SECRET_KEY=os.getenv("FLASK_SECRET", "dev-secret-change-me"),
    RATELIMIT_DEFAULT="200 per day;50 per hour",
    RATELIMIT_HEADERS_ENABLED=True
)

# Initialize rate limiter
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=[app.config["RATELIMIT_DEFAULT"]]
)

@app.errorhandler(429)
def ratelimit_handler(e):
    return jsonify({"error": "Rate limit exceeded. Please try again later."}), 429

@app.errorhandler(400)
def bad_request_error(e):
    app.logger.error(f"Bad Request Error: {str(e)}")
    flash("Bad request. Please try again.", "error")
    return redirect(request.referrer or url_for("index"))

# Helpers
def get_session_refs():
    refs = session.get("refs")
    if refs is None:
        # Try to pre-populate from persisted refs on first access
        refs = load_persisted_refs()
        session["refs"] = refs
    return refs

from flask_wtf import FlaskForm
from wtforms import StringField, SelectField, SubmitField, IntegerField, BooleanField
from wtforms.validators import DataRequired, Optional, NumberRange
from datetime import datetime

class SearchForm(FlaskForm):
    # Basic search fields
    query = StringField('Query', validators=[DataRequired()], 
                       render_kw={"placeholder": "Search by title, author, or keywords"})
    stype = SelectField('Type', choices=[
        ('title', 'Title / keywords'), 
        ('author', 'Author'),
        ('doi', 'DOI')
    ])
    style = SelectField('Style', choices=[
        ('harvard', 'Harvard'), 
        ('apa', 'APA'), 
        ('ieee', 'IEEE'),
        ('mla', 'MLA'),
        ('chicago', 'Chicago'),
        ('vancouver', 'Vancouver')
    ])
    
    # Advanced search fields
    year_from = IntegerField('From', validators=[
        Optional(), 
        NumberRange(min=1900, max=datetime.now().year)
    ], render_kw={"placeholder": "Year from"})
    
    year_to = IntegerField('To', validators=[
        Optional(), 
        NumberRange(min=1900, max=datetime.now().year)
    ], render_kw={"placeholder": "Year to"})
    
    document_type = SelectField('Document Type', choices=[
        ('', 'All Types'),
        ('article', 'Journal Article'),
        ('book', 'Book'),
        ('conference', 'Conference Paper'),
        ('thesis', 'Thesis'),
        ('report', 'Report')
    ], default='')
    
    language = SelectField('Language', choices=[
        ('', 'Any Language'),
        ('en', 'English'),
        ('fr', 'French'),
        ('de', 'German'),
        ('es', 'Spanish'),
        ('zh', 'Chinese')
    ], default='')
    
    include_abstract = BooleanField('Include abstract in search', default=False)
    open_access = BooleanField('Open access only', default=False)
    
    submit = SubmitField('Search', render_kw={"class": "btn btn-primary"})
    reset = SubmitField('Reset', render_kw={"class": "btn btn-outline-secondary"})

@app.route("/", methods=["GET", "POST"])
@limiter.limit("10 per minute")  # Allow 10 requests per minute
@limiter.limit("200 per day")   # And 200 requests per day
def index():
    form = SearchForm()
    refs = get_session_refs()
    style = session.get("style", "harvard")
    
    # Handle form submission
    if form.validate_on_submit():
        try:
            query = form.query.data.strip()
            if not query:
                flash('Please enter a search term', 'error')
                return redirect(url_for('index'))
                
            search_type = form.stype.data
            style = form.style.data
            session["style"] = style  # Save the selected style
            
            # Get advanced search parameters
            year_from = form.year_from.data
            year_to = form.year_to.data
            document_type = form.document_type.data
            language = form.language.data
            include_abstract = form.include_abstract.data
            open_access = form.open_access.data
            
            # Log search parameters
            app.logger.info(
                f"Searching for '{query}' (type: {search_type}, "
                f"years: {year_from or '*'}-{year_to or '*'}, "
                f"type: {document_type or 'any'}, language: {language or 'any'})"
            )
            
            cache = referencing.load_cache()
            results = []
            
            # Build search parameters
            search_params = {
                'query': query,
                'search_type': search_type,
                'year_from': year_from,
                'year_to': year_to,
                'document_type': document_type,
                'language': language,
                'include_abstract': include_abstract,
                'open_access': open_access
            }
            
            if search_type == "author":
                # For author search, we only need the author name
                results = referencing.lookup_author_works(query, cache) or []
                app.logger.info(f"Found {len(results)} results for author search")
            else:
                # For title search, we can add the query to the search parameters
                search_params['query'] = query
                m = referencing.lookup_single_work(query, cache)
                if m:
                    # Apply additional filters to the result
                    filtered = True
                    if year_from and int(m.get('year', 0)) < year_from:
                        filtered = False
                    if year_to and int(m.get('year', 9999)) > year_to:
                        filtered = False
                    if document_type and m.get('pub_type', '').lower() != document_type.lower():
                        filtered = False
                    
                    if filtered:
                        results = [m]
                        app.logger.info(f"Found result for {search_type} search: {m.get('title', 'Unknown title')}")
                    else:
                        app.logger.info("Result filtered out based on search criteria")
                        results = []
                else:
                    app.logger.warning(f"No results found for query: {query}")
                    results = []
            
            # If it's an AJAX request, return JSON
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({
                    'success': True,
                    'results': results,
                    'count': len(results)
                })
            
            return render_template("results.html", 
                                query=query, 
                                results=results, 
                                search_type=search_type,
                                style=style,
                                form=form)  # Pass form back to template
                                
        except Exception as e:
            app.logger.error(f"Error during search: {str(e)}", exc_info=True)
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({
                    'success': False,
                    'error': str(e)
                }), 500
            flash('An error occurred while searching. Please try again.', 'error')
            return redirect(url_for('index'))
    
    # For GET requests or invalid form
    form.style.default = style
    form.process()  # Ensure the form is properly initialized
    
    # Add current year for the year range fields
    current_year = datetime.now().year
    
    return render_template("index.html", 
                         refs=refs, 
                         style=style, 
                         form=form,
                         current_year=current_year)


class ManualForm(FlaskForm):
    title = StringField('Title', validators=[DataRequired()])
    authors = StringField('Authors (comma-separated, e.g. "Smith, J., Doe, A.")')
    year = StringField('Year')
    journal = StringField('Journal / Publisher')
    publisher = StringField('Publisher (for books)')
    volume = StringField('Volume')
    issue = StringField('Issue')
    pages = StringField('Pages')
    doi = StringField('DOI')
    pub_type = SelectField('Type', choices=[
        ('book', 'Book'), 
        ('journal-article', 'Journal Article'),
        ('proceedings-article', 'Proceedings Article')
    ])
    submit = SubmitField('Add Reference')

@app.route('/manual', methods=['GET'])
def manual():
    """Render the manual entry form."""
    form = ManualForm()
    return render_template('manual.html', form=form)


@app.route('/manual_add', methods=['POST'])
def manual_add():
    """Handle manual reference submission and add to session references."""
    title = request.form.get('title', '').strip()
    authors_raw = request.form.get('authors', '').strip()
    year = request.form.get('year', '').strip()
    journal = request.form.get('journal', '').strip()
    publisher = request.form.get('publisher', '').strip()
    volume = request.form.get('volume', '').strip()
    issue = request.form.get('issue', '').strip()
    pages = request.form.get('pages', '').strip()
    doi = request.form.get('doi', '').strip()
    pub_type = request.form.get('pub_type', 'book')

    if not title:
        flash('Title is required')
        return redirect(url_for('manual'))

    # parse authors: accept a variety of common formats
    authors = normalize_authors(authors_raw)
    # ensure authors format matches "Surname, Given" expectation if user entered "Given Surname"
    # we won't attempt to reformat here; assume user provides surname-first where possible

    pub = {
        'source': 'manual',
        'pub_type': pub_type,
        'authors': authors or [],
        'year': year or 'n.d.',
        'title': title,
        'journal': journal,
        'publisher': publisher,
        'location': '',
        'volume': volume,
        'issue': issue,
        'pages': pages,
        'doi': doi,
    }

    refs = get_session_refs()
    refs.append(pub)
    session['refs'] = refs
    # persist manual refs to disk
    save_persisted_refs(refs)
    flash('Manual reference added')
    return redirect(url_for('bibliography'))


@app.route('/remove/<int:idx>', methods=['GET', 'POST'])
def remove(idx):
    refs = get_session_refs()
    if 0 <= idx < len(refs):
        refs.pop(idx)
        session['refs'] = refs
        save_persisted_refs(refs)
        flash('Removed item')
    else:
        flash('Invalid index')
    return redirect(url_for('bibliography'))


@app.route('/edit/<int:idx>', methods=['GET', 'POST'])
def edit(idx):
    refs = get_session_refs()
    if request.method == 'GET':
        if 0 <= idx < len(refs):
            pub = refs[idx]
            return render_template('manual.html', edit=True, idx=idx, pub=pub)
        flash('Invalid index')
        return redirect(url_for('bibliography'))

    # POST: update
    if 0 <= idx < len(refs):
        title = request.form.get('title', '').strip()
        authors_raw = request.form.get('authors', '').strip()
        year = request.form.get('year', '').strip()
        journal = request.form.get('journal', '').strip()
        publisher = request.form.get('publisher', '').strip()
        volume = request.form.get('volume', '').strip()
        issue = request.form.get('issue', '').strip()
        pages = request.form.get('pages', '').strip()
        doi = request.form.get('doi', '').strip()
        pub_type = request.form.get('pub_type', 'book')

        if not title:
            flash('Title is required')
            return redirect(url_for('edit', idx=idx))

        authors = normalize_authors(authors_raw)
        pub = {
            'source': 'manual',
            'pub_type': pub_type,
            'authors': authors or [],
            'year': year or 'n.d.',
            'title': title,
            'journal': journal,
            'publisher': publisher,
            'location': '',
            'volume': volume,
            'issue': issue,
            'pages': pages,
            'doi': doi,
        }
        refs[idx] = pub
        session['refs'] = refs
        save_persisted_refs(refs)
        flash('Reference updated')
    else:
        flash('Invalid index')
    return redirect(url_for('bibliography'))

@app.route("/search", methods=["POST"])
@limiter.limit("5 per minute")  # Stricter limit on search
@limiter.limit("50 per day")   # And 50 per day
def search():
    q = request.form.get("query", "").strip()
    search_type = request.form.get("stype", "title")
    # capture selected style and persist in session
    style = request.form.get("style") or session.get("style", "harvard")
    session["style"] = style
    if not q:
        flash("Please enter a query.")
        return redirect(url_for("index"))

    cache = referencing.load_cache()
    results = []
    if search_type == "author":
        results = referencing.lookup_author_works(q, cache)
    else:
        m = referencing.lookup_single_work(q, cache)
        if m:
            results = [m]

    # normalize results for template
    return render_template(
        "results.html", 
        query=q, 
        results=results, 
        search_type=search_type
    )
@app.route('/export/<format>')
@limiter.limit("10 per minute")
def export(format):
    """Export references in the specified format."""
    idx = request.args.get('idx', None)
    if idx is not None:
        # Export single reference
        try:
            idx = int(idx)
            refs = get_session_refs()
            if 0 <= idx < len(refs):
                refs = [refs[idx]]
            else:
                return "Invalid reference index", 400
        except (ValueError, IndexError):
            return "Invalid reference index", 400
    else:
        # Export all references
        refs = get_session_refs()
        
    if not refs:
        return "No references to export", 400
        
    try:
        if format == 'bibtex':
            return Response(
                referencing.export_bibtex(refs),
                mimetype='text/plain',
                headers={'Content-Disposition': f'attachment;filename=reference_{idx}.bib' if idx is not None else 'references.bib'}
            )
        elif format == 'ris':
            return Response(
                referencing.export_ris(refs),
                mimetype='text/plain',
                headers={'Content-Disposition': f'attachment;filename=reference_{idx}.ris' if idx is not None else 'references.ris'}
            )
        elif format == 'json':
            return jsonify(refs[0] if idx is not None else [r for r in refs])
        elif format == 'csv':
            output = StringIO()
            writer = csv.writer(output)
            # Write header
            writer.writerow(['Title', 'Authors', 'Year', 'Journal', 'Volume', 'Issue', 'Pages', 'DOI', 'URL'])
            # Write data
            for r in refs:
                writer.writerow([
                    r.get('title', ''),
                    '; '.join(r.get('authors', [])) if isinstance(r.get('authors'), list) else r.get('authors', ''),
                    r.get('year', ''),
                    r.get('journal', ''),
                    r.get('volume', ''),
                    r.get('issue', ''),
                    r.get('pages', ''),
                    r.get('doi', ''),
                    r.get('url', '')
                ])
            return Response(
                output.getvalue(),
                mimetype='text/csv',
                headers={'Content-Disposition': f'attachment;filename=reference_{idx}.csv' if idx is not None else 'references.csv'}
            )
        else:
            return f"Unsupported format: {format}", 400
            
    except Exception as e:
        app.logger.error(f"Export error: {str(e)}")
        return f"Error exporting references: {str(e)}", 500

def _export_csv(refs):
    """Export references in CSV format."""
    # Get all possible field names
    fieldnames = [
        'title', 'authors', 'year', 'journal', 'publisher', 'volume', 
        'issue', 'pages', 'doi', 'url', 'abstract', 'language', 'type'
    ]
    
    def generate():
        data = io.StringIO()
        writer = csv.DictWriter(data, fieldnames=fieldnames)
        writer.writeheader()
        yield data.getvalue()
        data.seek(0)
        data.truncate(0)
        
        for ref in refs:
            try:
                # Convert dict to Reference object if needed
                if isinstance(ref, dict):
                    ref = Reference.from_dict(ref)
                writer.writerow(ref.to_csv_row())
                yield data.getvalue()
                data.seek(0)
                data.truncate(0)
            except Exception as e:
                app.logger.error(f"Error converting reference to CSV: {str(e)}")
                continue
    
    return Response(
        generate(),
        mimetype="text/csv",
        headers={"Content-disposition": "attachment; filename=references.csv"}
    )

def _export_json(refs):
    """Export references in JSON format."""
    refs_list = []
    for ref in refs:
        try:
            # Convert dict to Reference object if needed
            if isinstance(ref, dict):
                ref = Reference.from_dict(ref)
            refs_list.append(ref.to_dict())
        except Exception as e:
            app.logger.error(f"Error converting reference to JSON: {str(e)}")
            continue
    
    return Response(
        json.dumps(refs_list, indent=2, ensure_ascii=False),
        mimetype="application/json",
        headers={"Content-disposition": "attachment; filename=references.json"}
    )

def _export_word(refs):
    """Export references in Word format."""
    try:
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        doc.add_heading('References', 0)
        
        style = session.get('style', 'apa')
        
        for ref in refs:
            try:
                # Convert dict to Reference object if needed
                if isinstance(ref, dict):
                    ref = Reference.from_dict(ref)
                
                # Add reference to document
                p = doc.add_paragraph()
                p.add_run(ref.format_citation(style)).italic = True
                p.paragraph_format.space_after = Pt(12)
            except Exception as e:
                app.logger.error(f"Error adding reference to Word doc: {str(e)}")
                continue
        
        # Save to a BytesIO object
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        
        return send_file(
            f,
            as_attachment=True,
            download_name='references.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except ImportError:
        flash('Word export requires python-docx package', 'error')
        return redirect(url_for('index'))  

@app.route("/add", methods=["POST"])
@limiter.limit("5 per minute")  # Stricter limit on reference addition
@limiter.limit("50 per day")   # And 50 per day
def add():
    app.logger.info(f"Add route called. Form data: {request.form.keys()}")
        
    # expects a JSON payload of the publication dict
    data = request.form.get("pub")
    if not data:
        flash("No publication data received", "error")
        return redirect(request.referrer or url_for("index"))
        
    try:
        pub = json.loads(data)
    except Exception as e:
        app.logger.error(f"Error parsing publication data: {str(e)}")
        flash("Invalid publication data format", "error")
        return redirect(request.referrer or url_for("index"))

    try:
        refs = get_session_refs()
        refs.append(pub)
        session["refs"] = refs
        # persist
        save_persisted_refs(refs)
        flash('Added to bibliography', 'success')
    except Exception as e:
        app.logger.error(f"Error adding reference: {str(e)}")
        flash("An error occurred while adding the reference", "error")
        
    return redirect(request.referrer or url_for("index"))


@app.route('/cite', methods=['POST'])
@limiter.limit("10 per minute")
@limiter.limit("100 per day")
def cite():
    """
    Generate in-text citation and full reference for a single publication.

    Expects form fields:
      - pub: JSON of the publication
      - style: citation style ('harvard', 'apa', 'ieee') optional
    """
    app.logger.info(f"Cite route called. Form data: {request.form.keys()}")
            
    pub_data = request.form.get("pub")
    style = request.form.get("style", "harvard")

    if not pub_data:
        flash("No publication data provided", "error")
        return redirect(request.referrer or url_for("index"))

    try:
        pub = json.loads(pub_data)
    except json.JSONDecodeError as e:
        app.logger.error(f"Error parsing publication data: {str(e)}")
        flash("Invalid publication data format", "error")
        return redirect(request.referrer or url_for("index"))

    try:
        # Generate the citation
        citation = referencing.in_text_citation(pub, style)
        full_ref = referencing.reference_entry(pub, style)
        
        # Flash the citation to the user
        flash(f"In-text citation ({style.upper()}): {citation}", "success")
        flash(f"Full reference: {full_ref}", "info")
        
        return redirect(request.referrer or url_for("index"))
        
    except Exception as e:
        app.logger.error(f"Error in cite for pub '{pub.get('title', 'Unknown')}': {str(e)}", exc_info=True)
        flash(f"Could not generate citation for '{pub.get('title', 'this reference')}'. It may be missing required fields.", "error")
        return redirect(request.referrer or url_for("index"))

@app.route("/bibliography", methods=["GET"])
@limiter.limit("10 per minute")  # Allow 10 requests per minute
@limiter.limit("200 per day")   # And 200 requests per day
def bibliography():
    refs = get_session_refs()
    style = request.args.get("style", session.get("style", "harvard"))
    session["style"] = style
    uniq = referencing.dedupe(refs)
    refs_sorted = referencing.sort_for_bibliography(uniq, style)
    return render_template("bibliography.html", refs=refs_sorted, style=style)

@app.route("/export_word", methods=["GET"])
@limiter.limit("10 per hour")  # Limit exports to prevent abuse
def export_word():
    idx = request.args.get('idx', None)
    if idx is not None:
        # Export single reference
        try:
            idx = int(idx)
            refs = get_session_refs()
            if 0 <= idx < len(refs):
                refs = [refs[idx]]
            else:
                return "Invalid reference index", 400
        except (ValueError, IndexError):
            return "Invalid reference index", 400
    else:
        # Export all references
        refs = get_session_refs()
    style = session.get("style", "harvard")
    uniq = referencing.dedupe(refs)
    refs_sorted = referencing.sort_for_bibliography(uniq, style)
    path = referencing.save_references_to_word(refs_sorted, referencing.DOWNLOAD_FOLDER, referencing.WORD_FILENAME, set(), style)
    # Return file to user
    try:
        return send_file(path, as_attachment=True)
    except Exception:
        flash(f"Saved to {path} (file download not available)")
        return redirect(url_for("bibliography"))

# Simple health route
@app.route("/health", methods=["GET"])
def health():
    return "ok", 200

# Error handlers
@app.errorhandler(404)
def not_found_error(error):
    return render_template('error.html', error='Page not found', error_code=404), 404

@app.errorhandler(500)
def internal_error(error):
    return render_template('error.html', error='Internal server error', error_code=500), 500

@app.errorhandler(429)
def ratelimit_handler(e):
    return handle_rate_limit(e)

# Rate limiting configuration
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"]
)

# Apply rate limiting to specific routes
limiter.limit("10 per minute")(search)
limiter.limit("5 per minute")(add)
limiter.limit("5 per minute")(cite)
limiter.limit("2 per minute")(export_word)

if __name__ == "__main__":
    # In production, use a production WSGI server like Gunicorn
    app.run(debug=os.environ.get('FLASK_DEBUG', 'False').lower() == 'true')
